import sys
import subprocess
import os

def install_and_import(package, import_name=None):
    if import_name is None:
        import_name = package
    try:
        __import__(import_name)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def parse_docx(file_path):
    install_and_import("python-docx", "docx")
    import docx
    
    print(f"--- Đọc nội dung từ {file_path} ---")
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            
    # Xử lý bảng nếu có
    for idx, table in enumerate(doc.tables):
        print(f"\n[Bảng {idx+1}]")
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            print(" | ".join(row_data))

def parse_excel(file_path):
    install_and_import("pandas")
    import pandas as pd
    
    print(f"--- Đọc nội dung từ {file_path} ---")
    try:
        # Đọc tất cả các sheet
        sheets = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, df in sheets.items():
            print(f"\n[Sheet: {sheet_name}]")
            print(df.to_markdown(index=False))
    except Exception as e:
        print(f"Lỗi khi đọc file Excel: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Sử dụng: python parse_input.py <đường_dẫn_file>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"File không tồn tại: {file_path}")
        sys.exit(1)
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".docx":
        parse_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        parse_excel(file_path)
    else:
        print("Định dạng file không được hỗ trợ (chỉ hỗ trợ .docx, .xlsx, .xls)")
        # Cố đọc như file text
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                print(f.read())
        except:
            pass
